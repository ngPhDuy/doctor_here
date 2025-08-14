import os
import time
import logging
from dotenv import load_dotenv

from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from langchain.schema import Document as LangChainDocument 
from docx import Document as DocxDocument
import requests

# ========== LOGGING ==========
logging.basicConfig(level=logging.INFO)
load_dotenv()

# ========== MODEL CONFIG ==========
LLAMA_MODEL_PATH = "models/t-llama-2-7b.Q4_K.gguf"
EMBEDDING_MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
VECTORSTORE_PATH = os.path.join(BASE_DIR, "vectorstore", "faiss_index")
DATA_DIR = os.path.join(BASE_DIR, "med_data")

# ========== LOAD LLAMA ==========
# try:
#     logging.info(f"🔁 Loading GGUF LLAMA model from {LLAMA_MODEL_PATH}...")
#     start = time.time()

#     model = Llama(
#         model_path=LLAMA_MODEL_PATH,
#         n_ctx=2048,
#         n_threads=multiprocessing.cpu_count(),  
#         n_gpu_layers=0,                    
#         verbose=False,
#     )

#     logging.info(f"✅ GGUF LLAMA model loaded in {time.time() - start:.2f}s")
# except Exception as e:
#     logging.error(f"❌ Failed to load GGUF LLAMA model: {e}")
#     raise

# ========== LOAD VECTORSTORE + EMBEDDING ==========
try:
    logging.info("🔁 Loading embedding model and FAISS vectorstore...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    vectorstore = FAISS.load_local(VECTORSTORE_PATH, embeddings, allow_dangerous_deserialization=True)

    ######### HYBRID RETRIEVER ######## 

    def load_docx_documents(data_dir):
        docs = []
        for filename in os.listdir(data_dir):
            if filename.endswith(".docx"):
                doc_path = os.path.join(data_dir, filename)
                doc = DocxDocument(doc_path)
                full_text = ""
                for para in doc.paragraphs:
                    full_text += para.text + "\n"
                docs.append(LangChainDocument(page_content=full_text))
        return docs

    docs = load_docx_documents(DATA_DIR)
    docs = [doc for doc in docs if doc.page_content.strip()]
    print(f"📄 Loaded {len(docs)} documents.")

    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = splitter.split_documents(docs)
    split_docs = [doc for doc in split_docs if doc.page_content.strip()]

    dense_retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 5})

    bm25_retriever = BM25Retriever.from_documents(split_docs)
    bm25_retriever.k = 5

    hybrid_retriever = EnsembleRetriever(
        retrievers=[dense_retriever, bm25_retriever],
        weights=[0.5, 0.5],  # Try [0.7, 0.3] if semantic is better
    )
    logging.info("✅ Vectorstore loaded and retriever initialized")
except Exception as e:
    logging.error(f"❌ Failed to load vectorstore: {e}")
    raise

# ========== BOT RESPONSE WITH RAG ==========

def ask_tllama(prompt: str, max_tokens: int = 512) -> str:

    payload = {
        "model": "gemma3:12b",
        "prompt": prompt,
        "stream": False
    }

    public_url = os.getenv("GEMMA_API_URL") 
    
    response = requests.post(f"{public_url}/api/generate", json=payload)
    response.raise_for_status() 

    data = response.json()

    return data.get("response", "") 

def get_bot_response(user_input: str) -> str:
    try:
        # 1. Retrieve related documents

        ###### HYBRID RETRIEVER ###### 

        retrieved_docs = hybrid_retriever.invoke(user_input)
        top_k_docs = retrieved_docs[:5]
        context = "\n\n".join([doc.page_content.strip()[:1000] for doc in top_k_docs])

        # 2. Compose prompt
        prompt = f"""Bạn là một AI Assistant hỗ trợ người dùng sử dụng ứng dụng chăm sóc sức khỏe "Doctor Here". Hãy trả lời ngắn gọn, rõ ràng, chính xác và thân thiện. Nếu người dùng hỏi về tính năng, hướng dẫn sử dụng hoặc thắc mắc liên quan đến bác sĩ, lịch hẹn, uống thuốc, nhắc nhở, kết quả khám... thì hãy trả lời đúng theo tài liệu hệ thống.

❗Không tự suy đoán thông tin ngoài tài liệu. Nếu không chắc chắn, hãy trả lời: "Xin lỗi, tôi chưa có thông tin về nội dung này."

Cấu trúc câu trả lời nên hướng dẫn từng bước nếu liên quan thao tác, hoặc mô tả ngắn nếu là giải thích.

Nếu người dùng đặt các câu hỏi mang tính Có/Không (Yes/No) thì hãy trả lời kèm theo giải thích ngắn gọn.

Dưới đây là ví dụ một số câu hỏi người dùng có thể hỏi bạn:

1. Ứng dụng Doctor Here dùng để làm gì?
2. Làm sao để tôi đặt lịch khám với bác sĩ?
3. Tôi có thể chia sẻ thông tin khám bệnh cho người thân được không?
4. Tính năng nhắc uống thuốc hoạt động như thế nào?
5. Tôi muốn xem lại kết quả khám bệnh thì vào đâu?
6. Có thể nhắn tin với bác sĩ không?
7. Nếu hết thời gian tư vấn, tôi có thể gia hạn được không?

Thông tin trong tài liệu hướng dẫn đã bao gồm toàn bộ luồng thao tác người dùng, vì vậy hãy trả lời chính xác theo nội dung tài liệu này.

Khi sẵn sàng, hãy chờ người dùng đặt câu hỏi."

[Ngữ cảnh]
{context}

[Câu hỏi]
{user_input}

[Trả lời]"""
        return prompt

        # # 3. Generate response using T-LLAMA GGUF model
        # logging.info("🧠 Generating response from GGUF model...")
        # output = model(
        #     prompt,
        #     max_tokens=128,
        #     temperature=0.1,
        #     top_k=50,
        #     top_p=0.95,
        #     stop=["</s>", "###"],
        # )
        # response = output["choices"][0]["text"].strip()
        # return response

    except Exception as e:
        logging.error(f"❌ Error during RAG response: {e}")
        return "Xin lỗi, tôi không thể xử lý câu hỏi này ngay lúc này."