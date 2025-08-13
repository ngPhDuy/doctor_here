
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
import os
from docx import Document
from langchain.schema import Document as LangChainDocument 

DATA_DIR = "med_data"
VECTORSTORE_DIR = "vectorstore/faiss_index"

print("📥 Loading DOCX documents...")

def load_docx_documents(data_dir):
    docs = []
    for filename in os.listdir(data_dir):
        if filename.endswith(".docx"):
            doc_path = os.path.join(data_dir, filename)
            doc = Document(doc_path)
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            docs.append(LangChainDocument(page_content=full_text))
    return docs

docs = load_docx_documents(DATA_DIR)
docs = [doc for doc in docs if doc.page_content.strip()]
print(f"📄 Loaded {len(docs)} documents.")

splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
split_docs = splitter.split_documents(docs)
split_docs = [doc for doc in split_docs if doc.page_content.strip()]
print(f"✂️ Split into {len(split_docs)} chunks.")
for i, chunk in enumerate(split_docs):
    print(f"Chunk {i+1} - Length: {len(chunk.page_content)}")

if not split_docs:
    print("❌ Không có đoạn văn bản nào để tạo vector.")
    exit()

print("🔎 Embedding and indexing...")
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
faiss_index = FAISS.from_documents(split_docs, embedding_model)

print(f"💾 Saving FAISS vectorstore to {VECTORSTORE_DIR}")
faiss_index.save_local(VECTORSTORE_DIR)

print("✅ Vectorstore created and saved successfully.")



